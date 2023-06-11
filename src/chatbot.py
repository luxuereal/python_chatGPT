from llama_index import SimpleDirectoryReader, GPTVectorStoreIndex, LLMPredictor, PromptHelper, StorageContext, load_index_from_storage, ServiceContext
from langchain.chat_models import ChatOpenAI
from docx import Document
import gradio as gr
import json
import os

os.environ["OPENAI_API_KEY"] = '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'

# max_input_size = 4096
# num_outputs = 512
# max_chunk_overlap = 20
# chunk_size_limit = 600

# prompt_helper = PromptHelper(
#     max_input_size=max_input_size, num_output=num_outputs, max_chunk_overlap=max_chunk_overlap, chunk_size_limit=chunk_size_limit)
# llm_predictor = LLMPredictor(llm=ChatOpenAI(
#     temperature=0.7, model_name="gpt-3.5-turbo", max_tokens=num_outputs))
# service_context = ServiceContext.from_defaults(
#     llm_predictor=llm_predictor, prompt_helper=prompt_helper)


def build_chatbot_ui():
    def get_answer(question):
        storage_context = StorageContext.from_defaults(
            persist_dir='./src/storage-all')
        # index = load_index_from_storage(
        #     storage_context, service_context=service_context)
        index = load_index_from_storage(storage_context)
        query_engine = index.as_query_engine()
        response = query_engine.query(question)

        with open("history.json", "a") as history_file:
            history_file.write(
                json.dumps({'Question': question, 'Answer': response.response}) + "\n")

        return response.response

    def load_history(chat_history):
        chat_history = []
        with open("history.json", "r") as history_file:
            lines = history_file.readlines()
            history_file.close()
        for line in lines:
            chat = json.loads(line)
            chat_history.append((chat["Question"], chat["Answer"]))
        return chat_history

    def save_to_docx():
        chat_history = []
        with open("history.json", "r") as history_file:
            lines = history_file.readlines()
            history_file.close()

        document = Document()
        for line in lines:
            chat = json.loads(line)
            document.add_paragraph("Question: " + chat["Question"])
            if chat["Answer"] is not None:
                document.add_paragraph("Answer: " + (chat["Answer"]))
            else:
                document.add_paragraph("Answer: No Answer.")
        document.save("history.docx")
        return chat_history

    with gr.Blocks(title="QA Chatbot", theme=gr.themes.Default(primary_hue="red", secondary_hue="pink")) as interface:
        with gr.Tab("QA"):
            with gr.Row():
                with gr.Column():
                    question = gr.Textbox(label="Question")
                with gr.Column():
                    answer = gr.Textbox(label="Answer", interactive=False)
            with gr.Row():
                submit = gr.Button("Submit", variant="primary")
        with gr.Tab("History"):
            chatbot = gr.Chatbot([], label="History").style(height=750)
            with gr.Row():
                with gr.Column():
                    reload = gr.Button("Load History", variant="primary")
                with gr.Column():
                    savetodoc = gr.Button("Publish")

        question.submit(fn=get_answer, inputs=question, outputs=answer)
        submit.click(fn=get_answer, inputs=question, outputs=answer)
        reload.click(fn=load_history, inputs=chatbot, outputs=chatbot)
        savetodoc.click(fn=save_to_docx)
    return interface


chatbot = build_chatbot_ui()
chatbot.launch(share=True)
